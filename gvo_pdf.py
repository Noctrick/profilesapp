import os
import pandas as pd
from docx import Document
from datetime import datetime
import subprocess
import shutil
import sys

# === CONFIGURATION ===
EXCEL_FILE = 'GVO_certificaat_template.xlsx'
TEMPLATE_DOCX = 'GVO_certificaat.docx'

# Optional filters
SPECIFIC_KLANT = None  # Example: 'Patrick B.V.'
SPECIFIC_EAN = None
SPECIFIC_ADRES = None

DOCX_OUTPUT_FOLDER = "Generated_Certificates"
PDF_OUTPUT_FOLDER = "Generated_PDFs"
ZIP_FILENAME = "Generated_PDFs.zip"

# === UTILITY FUNCTIONS ===

def replace_text_in_paragraph(paragraph, replacements):
    for run in paragraph.runs:
        for key, value in replacements.items():
            if key in run.text:
                run.text = run.text.replace(key, value)

def replace_text_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for key, value in replacements.items():
                if key in cell.text:
                    cell.text = cell.text.replace(key, value)

def convert_docx_to_pdf(input_docx, output_pdf):
    """ Convert a .docx file to .pdf using LibreOffice """
    command = [
        "libreoffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", os.path.dirname(output_pdf),
        input_docx
    ]
    subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    if os.path.exists(output_pdf):
        print(f"✅ PDF generated: {output_pdf}")
        return output_pdf
    else:
        print(f"❌ Failed to convert {input_docx} to PDF.")
        return None

# === MAIN LOGIC ===

def main():
    # Create output folders
    os.makedirs(DOCX_OUTPUT_FOLDER, exist_ok=True)
    os.makedirs(PDF_OUTPUT_FOLDER, exist_ok=True)

    # Load Excel data
    try:
        df = pd.read_excel(EXCEL_FILE, engine='openpyxl')
        df.columns = df.columns.str.strip()
    except Exception as e:
        print(f"❌ Failed to load Excel file: {e}")
        sys.exit(1)

    # Apply filters
    if SPECIFIC_ADRES:
        df = df[df['STRAAT'].str.contains(SPECIFIC_ADRES, na=False)]

    if SPECIFIC_KLANT:
        df = df[df['KLANTNAAM'] == SPECIFIC_KLANT]
    if SPECIFIC_EAN:
        df = df[df['EAN'] == SPECIFIC_EAN]

    if df.empty:
        print("⚠️ No matching entries found.")
        sys.exit()

    for index, row in df.iterrows():
        klant = str(row['KLANTNAAM'])
        ean = str(row['EAN'])
        leveringsadres1 = str(row['STRAAT'])
        leveringsadres2 = str(row['STAD'])
        datum_vandaag = datetime.today().strftime("%d-%m-%Y")

        try:
            doc = Document(TEMPLATE_DOCX)
        except Exception as e:
            print(f"❌ Could not load Word template: {e}")
            continue

        replacements = {
            "KOLOM A": klant,
            "KOLOM B": ean,
            "KOLOM C": leveringsadres1,
            "KOLOM D": leveringsadres2,
            "Datum: van vandaag": f"Datum: {datum_vandaag}"
        }

        for para in doc.paragraphs:
            replace_text_in_paragraph(para, replacements)

        for table in doc.tables:
            replace_text_in_table(table, replacements)

        safe_klant = klant.replace("/", "-").replace(" ", "_")
        safe_city = leveringsadres2.replace("/", "-").replace(" ", "_")
        safe_straat = leveringsadres1.replace("/", "-").replace(" ", "_")

        docx_filename = f'GVO_Certificaat_{safe_klant}_{safe_city}_{safe_straat}.docx'
        pdf_filename = f'GVO_Certificaat_{safe_klant}_{safe_city}_{safe_straat}.pdf'

        output_docx = os.path.join(DOCX_OUTPUT_FOLDER, docx_filename)
        output_pdf = os.path.join(PDF_OUTPUT_FOLDER, pdf_filename)

        doc.save(output_docx)
        print(f"📄 Certificate saved: {output_docx}")

        converted_pdf = convert_docx_to_pdf(output_docx, output_pdf)

        if converted_pdf and os.path.exists(output_docx):
            os.remove(output_docx)
            print(f"🧹 Deleted intermediate DOCX: {output_docx}")

    # Create ZIP of PDFs
    shutil.make_archive(ZIP_FILENAME.replace(".zip", ""), 'zip', PDF_OUTPUT_FOLDER)
    print(f"📦 ZIP file created: {ZIP_FILENAME}")


if __name__ == "__main__":
    main()
